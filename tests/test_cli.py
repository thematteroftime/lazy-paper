from pathlib import Path
from unittest.mock import patch

import pytest
import yaml

from cli import main


def test_cli_run_creates_run_dir(tmp_path: Path, monkeypatch):
    monkeypatch.setenv("PADDLEOCR_TOKEN", "fake")
    monkeypatch.setenv("LLM_VISION_API_KEY", "fake")
    monkeypatch.setenv("LLM_TEXT_API_KEY", "fake")

    pdf = tmp_path / "paper.pdf"
    pdf.write_bytes(b"%PDF-1.4\n")
    tpl = tmp_path / "tpl.docx"
    from docx import Document
    d = Document(); d.add_paragraph("Outline"); d.save(tpl)

    runs_dir = tmp_path / "runs"

    def mk_runner(name):
        def fake_run(**kwargs):
            outd = kwargs.get("out_dir") or Path("/tmp")
            outd.mkdir(parents=True, exist_ok=True)
            (outd / "done.yaml").write_text("ok\n", encoding="utf-8")
            return {"name": name}
        return fake_run

    targets = [
        "stages.s01_ocr.runner.run",
        "stages.s02_clean.runner.run",
        "stages.s03_chapter.runner.run",
        "stages.s04_figures.runner.run",
        "stages.s05_template.runner.run",
        "stages.s06_context.runner.run",
        "stages.s07_figure_analyze.runner.run",
        "stages.s08_section_compose.runner.run",
        "stages.s09_render.runner.run",
    ]
    patches = [patch(t, mk_runner(t)) for t in targets]
    for pp in patches:
        pp.start()
    try:
        rc = main([
            "run",
            "--pdf", str(pdf),
            "--template", str(tpl),
            "--runs-dir", str(runs_dir),
            "--paper-id", "paper",
        ])
    finally:
        for pp in patches:
            pp.stop()

    assert rc == 0
    meta = yaml.safe_load((runs_dir / "paper" / "meta.yaml").read_text(encoding="utf-8"))
    assert meta["paper_id"] == "paper"
    assert meta["stages_completed"] == [t.rsplit(".", 2)[0].split(".")[1] for t in targets]


def test_cli_passes_formats_to_s09(tmp_path: Path, monkeypatch):
    monkeypatch.setenv("PADDLEOCR_TOKEN", "fake")
    monkeypatch.setenv("LLM_VISION_API_KEY", "fake")
    monkeypatch.setenv("LLM_TEXT_API_KEY", "fake")

    pdf = tmp_path / "p.pdf"; pdf.write_bytes(b"%PDF-1.4\n")
    tpl = tmp_path / "t.docx"
    from docx import Document as DocxDocument
    DocxDocument().save(tpl)

    captured: dict = {}

    def mk_runner(name):
        def fake_run(**kwargs):
            outd = kwargs["out_dir"]; outd.mkdir(parents=True, exist_ok=True)
            (outd / "done.yaml").write_text("ok\n", encoding="utf-8")
            if name == "stages.s09_render.runner.run":
                captured.update(kwargs)
            return {"name": name}
        return fake_run

    targets = [f"stages.{s}.runner.run" for s in [
        "s01_ocr", "s02_clean", "s03_chapter", "s04_figures", "s05_template",
        "s06_context", "s07_figure_analyze", "s08_section_compose", "s09_render",
    ]]
    patches = [patch(t, mk_runner(t)) for t in targets]
    for pp in patches: pp.start()
    try:
        from cli import main
        rc = main([
            "run", "--pdf", str(pdf), "--template", str(tpl),
            "--runs-dir", str(tmp_path / "runs"), "--paper-id", "p",
            "--formats", "docx,pptx", "--pptx-bullets", "rule",
        ])
    finally:
        for pp in patches: pp.stop()

    assert rc == 0
    assert captured.get("formats") == ["docx", "pptx"]
    assert captured.get("pptx_bullets") == "rule"


def test_cli_only_splits_comma_separated_stages(tmp_path: Path, monkeypatch):
    monkeypatch.setenv("PADDLEOCR_TOKEN", "fake")
    monkeypatch.setenv("LLM_VISION_API_KEY", "fake")
    monkeypatch.setenv("LLM_TEXT_API_KEY", "fake")

    pdf = tmp_path / "p.pdf"; pdf.write_bytes(b"%PDF-1.4\n")
    tpl = tmp_path / "t.docx"
    from docx import Document as DocxDocument
    DocxDocument().save(tpl)

    seen: list[str] = []
    targets = [f"stages.{s}.runner.run" for s in [
        "s01_ocr", "s02_clean", "s03_chapter", "s04_figures", "s05_template",
        "s06_context", "s07_figure_analyze", "s08_section_compose", "s09_render",
    ]]

    def mk_runner(stage_id):
        def fake_run(**kwargs):
            outd = kwargs["out_dir"]; outd.mkdir(parents=True, exist_ok=True)
            (outd / "done.yaml").write_text("ok\n", encoding="utf-8")
            seen.append(stage_id)
            return {}
        return fake_run

    patches = [patch(t, mk_runner(t.split(".")[1])) for t in targets]
    for pp in patches: pp.start()
    try:
        rc = main([
            "run", "--pdf", str(pdf), "--template", str(tpl),
            "--runs-dir", str(tmp_path / "runs"), "--paper-id", "p",
            "--only", "s08_section_compose,s09_render",
            "--force",
        ])
    finally:
        for pp in patches: pp.stop()

    assert rc == 0
    assert seen == ["s08_section_compose", "s09_render"]


def test_cli_only_rejects_unknown_stage(tmp_path: Path, monkeypatch):
    monkeypatch.setenv("PADDLEOCR_TOKEN", "fake")
    monkeypatch.setenv("LLM_VISION_API_KEY", "fake")
    monkeypatch.setenv("LLM_TEXT_API_KEY", "fake")

    pdf = tmp_path / "p.pdf"; pdf.write_bytes(b"%PDF-1.4\n")
    tpl = tmp_path / "t.docx"
    from docx import Document as DocxDocument
    DocxDocument().save(tpl)

    with pytest.raises(SystemExit, match="Unknown stage"):
        main([
            "run", "--pdf", str(pdf), "--template", str(tpl),
            "--runs-dir", str(tmp_path / "runs"), "--paper-id", "p",
            "--only", "s09_renderr",  # typo
        ])
