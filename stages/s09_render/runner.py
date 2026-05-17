"""Stage 09: render mypaper-compatible bundle + self-contained preview docx."""
from __future__ import annotations

import re
import shutil
from pathlib import Path

import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from stages._common import mark_done

BUNDLE_README = """\
# mypaper bundle

Drop this folder's contents into mypaper/ to render the styled thesis:

    cp -r chapters/* /path/to/mypaper/chapters/
    cp -r figures/*  /path/to/mypaper/figures/
    cd /path/to/mypaper && uv run python scripts/build.py

The README of mypaper has the full template-swap instructions.
"""


def _cn_font(run, size=10.5, bold=False, color=None, set_eastasia: bool = True):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    if set_eastasia:
        rPr = run._element.get_or_add_rPr()
        rf = rPr.find(qn("w:rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts")
            rPr.append(rf)
        rf.set(qn("w:eastAsia"), "宋体")
        rf.set(qn("w:ascii"), "Times New Roman")
        rf.set(qn("w:hAnsi"), "Times New Roman")


def _render_preview_docx(*, compose_dir: Path, fig_notes: list[dict],
                         out_path: Path, paper_title: str, lang: str = "zh") -> None:
    body_pt = 10.5 if lang == "zh" else 11
    img_cm = 13 if lang == "zh" else 14
    title_pt = 16
    heading_pt = 14
    caption_pt = 9
    set_ea = (lang == "zh")

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Cm(2.0)
    sec.left_margin = sec.right_margin = Cm(2.2)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _cn_font(p.add_run(paper_title), size=title_pt, bold=True, set_eastasia=set_ea)

    fig_by_id = {n["fig_id"]: n for n in fig_notes}
    embedded: set[str] = set()

    for ch in sorted((compose_dir / "chapters").glob("*.md")):
        text = ch.read_text(encoding="utf-8")
        lines = text.splitlines()
        i = 0
        if lines and lines[0].startswith("# "):
            heading = lines[0][2:].strip()
            p = doc.add_paragraph()
            _cn_font(p.add_run(heading), size=heading_pt, bold=True, set_eastasia=set_ea)
            i = 1
        body = "\n".join(lines[i:]).strip()
        for para in body.split("\n\n"):
            if not para.strip():
                continue
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0.74)
            _cn_font(p.add_run(para.strip()), size=body_pt, set_eastasia=set_ea)
        for fid, note in fig_by_id.items():
            if fid in embedded:
                continue
            # Match both English ("Fig. 1") and Chinese ("图1", "图 1") references
            m = re.match(r"Fig\.\s*(\d+)", fid)
            fig_num = m.group(1) if m else ""
            cn_refs = (f"图{fig_num}", f"图 {fig_num}") if fig_num else ()
            referenced = (fid in body) or any(ref in body for ref in cn_refs)
            if not referenced:
                continue
            # Collect all image paths (prefer image_paths list if present, else image_abs_path)
            paths = note.get("image_paths") or []
            if not paths and note.get("image_abs_path"):
                paths = [note["image_abs_path"]]
            paths = [p for p in paths if p and Path(p).exists()]
            if not paths:
                continue
            cap_text = note.get("caption") or note.get("caption_cn") or fid
            deep_obs = note.get("deep_observation") or note.get("deep_observation_cn") or ""
            # Render each panel as its own centered image (no extra caption per panel)
            for p in paths:
                ip = doc.add_paragraph(); ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
                ip.add_run().add_picture(p, width=Cm(img_cm))
            # One caption under the whole block
            cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            label = f"{fid}" if lang == "en" else f"{fid.replace('Fig.', '图')}"
            _cn_font(cap.add_run(f"{label}. {cap_text}"), size=caption_pt, bold=True, set_eastasia=set_ea)
            if deep_obs:
                obs = doc.add_paragraph()
                obs_prefix = "【深度观察】" if lang == "zh" else "Deep observation: "
                _cn_font(obs.add_run(f"{obs_prefix}{deep_obs}"),
                         size=caption_pt, color=(0x33, 0x33, 0x66), set_eastasia=set_ea)
            embedded.add(fid)
    doc.save(out_path)


def run(*, compose_dir: Path, fig_notes_dir: Path, out_dir: Path,
        paper_title: str = "Paper Preview", lang: str = "zh") -> dict:
    out_dir.mkdir(parents=True, exist_ok=True)
    bundle = out_dir / "mypaper_bundle"
    (bundle / "chapters").mkdir(parents=True, exist_ok=True)
    (bundle / "figures").mkdir(exist_ok=True)

    # Clear stale files from prior runs so we don't accumulate
    for stale in (bundle / "chapters").glob("*.md"):
        stale.unlink()
    for stale in (bundle / "figures").iterdir():
        if stale.is_file():
            stale.unlink()

    for md in (compose_dir / "chapters").glob("*.md"):
        shutil.copy2(md, bundle / "chapters" / md.name)

    fig_notes = yaml.safe_load((fig_notes_dir / "fig_notes.yaml").read_text(encoding="utf-8")) or []
    for note in fig_notes:
        # Copy ALL sub-panel images (image_paths) plus the canonical one (image_abs_path)
        paths = list(note.get("image_paths") or [])
        if note.get("image_abs_path"):
            paths.append(note["image_abs_path"])
        for p in paths:
            ap = Path(p)
            if ap.exists():
                shutil.copy2(ap, bundle / "figures" / ap.name)

    (bundle / "README.md").write_text(BUNDLE_README, encoding="utf-8")

    preview = out_dir / "preview.docx"
    _render_preview_docx(compose_dir=compose_dir, fig_notes=fig_notes,
                         out_path=preview, paper_title=paper_title, lang=lang)

    mark_done(out_dir, {
        "bundle_chapters": len(list((bundle / "chapters").glob("*.md"))),
        "bundle_figures": len(list((bundle / "figures").glob("*"))),
        "preview_bytes": preview.stat().st_size,
    })
    return {"preview": str(preview), "bundle": str(bundle)}
