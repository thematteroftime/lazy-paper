"""Render a Document to .docx via python-docx, sharing the HTML renderer's
design tokens (accent orange, serif headings, secondary-gray captions,
accent-bordered deep-observation block)."""
from __future__ import annotations

from pathlib import Path
from typing import ClassVar

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from stages.s09_render._math import iter_runs
from stages.s09_render.model import Chapter, Document, FigureBlock, Paragraph, TableBlock
from stages.s09_render.renderers import RENDERERS
from stages.s09_render.renderers.base import Renderer


# Hand-synced with the corresponding `:root` variables in styles.css.
_ACCENT_RGB = (0xD9, 0x77, 0x57)
_INK_PRIMARY_RGB = (0x1F, 0x1B, 0x16)
_INK_SECONDARY_RGB = (0x5E, 0x58, 0x51)


class DocxRenderer(Renderer):
    extension: ClassVar[str] = "docx"

    TITLE_PT = 18
    HEADING_PT = 14
    CHAPTER_NUM_PT = 11
    CAPTION_PT = 9

    def render(self, doc: Document, out_path: Path) -> None:
        body_pt = 10.5 if doc.lang == "zh" else 11
        img_cm = 13 if doc.lang == "zh" else 14
        set_ea = (doc.lang == "zh")

        out_doc = DocxDocument()
        sec = out_doc.sections[0]
        sec.top_margin = sec.bottom_margin = Cm(2.0)
        sec.left_margin = sec.right_margin = Cm(2.2)

        self._write_title(out_doc, doc.paper_title, set_ea)
        for idx, chapter in enumerate(doc.chapters, start=1):
            self._write_chapter(out_doc, chapter, idx, body_pt, img_cm, set_ea, doc.lang)

        out_doc.save(out_path)

    # ---------- chapter / block writers ----------

    def _write_chapter(self, out, chapter: Chapter, idx: int, body_pt: float,
                       img_cm: float, set_ea: bool, lang: str) -> None:
        self._write_heading(out, chapter.heading, idx, set_ea)
        for block in chapter.blocks:
            if isinstance(block, Paragraph):
                self._write_paragraph(out, block.text, body_pt, set_ea)
            elif isinstance(block, FigureBlock):
                self._write_figure_block(out, block, img_cm, set_ea, lang)
            elif isinstance(block, TableBlock):
                self._write_table_block(out, block, set_ea)

    def _write_title(self, out, title: str, set_ea: bool) -> None:
        p = out.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._apply_cn_font(
            p.add_run(title),
            size=self.TITLE_PT, bold=True, set_ea=set_ea,
            color=_INK_PRIMARY_RGB, serif=True,
        )

    def _write_heading(self, out, heading: str, idx: int, set_ea: bool) -> None:
        try:
            p = out.add_paragraph(style="Heading 1")
        except KeyError:
            p = out.add_paragraph()
        # Two runs: "01  " in accent mono, then heading text in serif bold.
        num_run = p.add_run(f"{idx:02d}  ")
        self._apply_mono_font(num_run, size=self.CHAPTER_NUM_PT, bold=True,
                              color=_ACCENT_RGB)
        head_run = p.add_run(heading)
        self._apply_cn_font(head_run, size=self.HEADING_PT, bold=True,
                            set_ea=set_ea, color=_INK_PRIMARY_RGB, serif=True)
        # Accent left border, mimicking the HTML chapter-heading::before bar.
        self._add_left_border(p, _ACCENT_RGB, size_eighths_pt=18)

    def _write_paragraph(self, out, text: str, body_pt: float, set_ea: bool) -> None:
        p = out.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        # Split into styled runs so **bold** and inline math (wrapped by
        # normalize_math) get bold / italic respectively. Plain prose remains
        # unchanged. Citation `[span:..]` markers in the text are still
        # processed by self._process_text on each segment.
        for segment, style in iter_runs(self._process_text(text)):
            if not segment:
                continue
            run = p.add_run(segment)
            self._apply_cn_font(run, size=body_pt, set_ea=set_ea,
                                bold=(style == "bold"),
                                italic=(style == "italic"))

    def _write_figure_block(self, out, block: FigureBlock,
                            img_cm: float, set_ea: bool, lang: str) -> None:
        paths = [p for p in block.image_paths if p.exists()]
        if not paths:
            return
        for img_path in paths:
            ip = out.add_paragraph()
            ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ip.add_run().add_picture(str(img_path), width=Cm(img_cm))
        cap = out.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # "图 5." in accent, " caption…" in secondary ink.
        tag_run = cap.add_run(f"{block.label}. ")
        self._apply_cn_font(tag_run, size=self.CAPTION_PT, bold=True,
                            set_ea=set_ea, color=_ACCENT_RGB)
        body_run = cap.add_run(block.caption)
        self._apply_cn_font(body_run, size=self.CAPTION_PT, bold=True,
                            set_ea=set_ea, color=_INK_SECONDARY_RGB)
        if block.deep_observation:
            self._write_deep_obs(out, block.deep_observation, set_ea, lang)

    def _write_deep_obs(self, out, text: str, set_ea: bool, lang: str) -> None:
        p = out.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        prefix = "⌖ 深度观察 " if lang == "zh" else "⌖ Deep observation "
        label_run = p.add_run(prefix)
        self._apply_cn_font(label_run, size=self.CAPTION_PT, bold=True,
                            set_ea=set_ea, color=_ACCENT_RGB)
        body_run = p.add_run(text)
        self._apply_cn_font(body_run, size=self.CAPTION_PT, italic=True,
                            set_ea=set_ea, color=_INK_SECONDARY_RGB)
        self._add_left_border(p, _ACCENT_RGB, size_eighths_pt=24)

    def _write_table_block(self, out, block: TableBlock, set_ea: bool) -> None:
        n_cols = len(block.headers)
        if n_cols == 0:
            return
        n_rows = 1 + len(block.rows)
        try:
            table = out.add_table(rows=n_rows, cols=n_cols, style="Light Grid")
        except KeyError:
            table = out.add_table(rows=n_rows, cols=n_cols)
        # Header row
        for j, h in enumerate(block.headers):
            cell = table.rows[0].cells[j]
            cell.text = h
            if cell.paragraphs:
                run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(h)
                run.bold = True
        # Data rows
        for i, row in enumerate(block.rows):
            for j, c in enumerate(row):
                if j < n_cols:
                    table.rows[i + 1].cells[j].text = c

    # ---------- font + XML helpers ----------

    @staticmethod
    def _apply_cn_font(run, *, size: float, bold: bool = False,
                       italic: bool = False,
                       color: tuple[int, int, int] | None = None,
                       set_ea: bool = True, serif: bool = False) -> None:
        # serif=True swaps the CJK face to Songti for title / heading runs.
        latin = "Times New Roman"
        eastasia = "宋体" if not serif else "Songti SC"
        run.font.name = latin
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = RGBColor(*color)
        if set_ea:
            rPr = run._element.get_or_add_rPr()
            rf = rPr.find(qn("w:rFonts"))
            if rf is None:
                rf = OxmlElement("w:rFonts")
                rPr.append(rf)
            rf.set(qn("w:eastAsia"), eastasia)
            rf.set(qn("w:ascii"), latin)
            rf.set(qn("w:hAnsi"), latin)

    @staticmethod
    def _apply_mono_font(run, *, size: float, bold: bool = False,
                         color: tuple[int, int, int] | None = None) -> None:
        run.font.name = "Menlo"
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
        # Override eastAsia hint so the mono ch-num doesn't fall back to
        # 宋体 when adjacent to CJK text.
        rPr = run._element.get_or_add_rPr()
        rf = rPr.find(qn("w:rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts")
            rPr.append(rf)
        for attr in ("w:eastAsia", "w:ascii", "w:hAnsi"):
            rf.set(qn(attr), "Menlo")

    @staticmethod
    def _add_left_border(p, rgb: tuple[int, int, int], *,
                         size_eighths_pt: int = 24) -> None:
        """Add a left-side colored border on a paragraph (mimics the HTML
        ``::before`` accent bar). size is in eighths of a point, e.g. 24 → 3pt.
        """
        pPr = p._p.get_or_add_pPr()
        pBdr = pPr.find(qn("w:pBdr"))
        if pBdr is None:
            pBdr = OxmlElement("w:pBdr")
            pPr.append(pBdr)
        left = pBdr.find(qn("w:left"))
        if left is None:
            left = OxmlElement("w:left")
            pBdr.append(left)
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), str(size_eighths_pt))
        left.set(qn("w:space"), "6")
        left.set(qn("w:color"), f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")


RENDERERS["docx"] = DocxRenderer
