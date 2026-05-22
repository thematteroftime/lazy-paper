"""Tests: citation marker handling in DOCX and HTML renderers."""
from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document as DocxDocument

from llm.citation import CitationMode
from stages.s09_render.model import Chapter, Document, Paragraph
from stages.s09_render.renderers import RENDERERS
import stages.s09_render.renderers.docx  # noqa: F401
import stages.s09_render.renderers.html  # noqa: F401

MARKER_TEXT = "Some text [span:doc123:10-20] with a citation."
CLEAN_TEXT = "Some text  with a citation."


def _make_doc() -> Document:
    return Document(
        paper_title="Citation Test",
        lang="en",
        chapters=(
            Chapter(heading="Intro", level=1, blocks=(
                Paragraph(text=MARKER_TEXT),
            )),
        ),
    )


def test_docx_remove_strips_markers(tmp_path: Path):
    """Default REMOVE mode: [span:…] markers are absent in the rendered DOCX."""
    doc = _make_doc()
    out = tmp_path / "preview.docx"
    RENDERERS["docx"]().render(doc, out)
    d = DocxDocument(out)
    full_text = "\n".join(p.text for p in d.paragraphs)
    assert "[span:" not in full_text
    assert "Some text" in full_text


def test_docx_keep_preserves_markers(tmp_path: Path):
    """KEEP mode: [span:…] markers are preserved verbatim in the rendered DOCX."""
    doc = _make_doc()
    out = tmp_path / "preview.docx"
    RENDERERS["docx"](citation_mode=CitationMode.KEEP).render(doc, out)
    d = DocxDocument(out)
    full_text = "\n".join(p.text for p in d.paragraphs)
    assert "[span:doc123:10-20]" in full_text


def test_html_remove_strips_markers(tmp_path: Path, monkeypatch):
    """REMOVE mode: [span:…] markers removed AND no <sup> anchor inserted.

    The old test only checked `[span:` absence — but HYPERLINK mode also
    satisfies that (it rewrites markers into <sup><a> anchors). Asserting
    `<sup class="cite-anchor"` absence discriminates REMOVE from HYPERLINK
    (the CSS always references the class name even when no anchors
    render, so we match the literal HTML tag).
    """
    monkeypatch.delenv("LAZY_PAPER_HTML_CITATIONS", raising=False)
    doc = _make_doc()
    out = tmp_path / "preview.html"
    RENDERERS["html"](citation_mode=CitationMode.REMOVE).render(doc, out)
    html = out.read_text(encoding="utf-8")
    assert "[span:" not in html
    assert '<sup class="cite-anchor"' not in html  # discriminates from HYPERLINK
    assert "Some text" in html


def test_html_keep_preserves_markers(tmp_path: Path):
    """KEEP mode: [span:…] markers are preserved verbatim in the rendered HTML."""
    doc = _make_doc()
    out = tmp_path / "preview.html"
    RENDERERS["html"](citation_mode=CitationMode.KEEP).render(doc, out)
    html = out.read_text(encoding="utf-8")
    assert "[span:doc123:10-20]" in html


def test_html_hyperlink_emits_anchor_and_sources_footer(tmp_path: Path, monkeypatch):
    """HYPERLINK mode (the new default): markers become clickable anchors
    and a sources-footer section is emitted."""
    monkeypatch.delenv("LAZY_PAPER_HTML_CITATIONS", raising=False)
    doc = _make_doc()
    out = tmp_path / "preview.html"
    RENDERERS["html"](citation_mode=CitationMode.HYPERLINK).render(doc, out)
    html = out.read_text(encoding="utf-8")
    assert "[span:" not in html  # raw marker stripped
    assert '<sup class="cite-anchor"' in html  # anchor emitted
    assert 'id="sources"' in html  # footer present (CSS-class name is "sources-footer", but the section id is "sources")
    assert "Some text" in html


def test_html_env_override_remove(tmp_path: Path, monkeypatch):
    """LAZY_PAPER_HTML_CITATIONS=remove overrides the default HYPERLINK."""
    monkeypatch.setenv("LAZY_PAPER_HTML_CITATIONS", "remove")
    doc = _make_doc()
    out = tmp_path / "preview.html"
    # Caller passes default HYPERLINK — env should win.
    RENDERERS["html"]().render(doc, out)
    html = out.read_text(encoding="utf-8")
    assert "[span:" not in html
    assert '<sup class="cite-anchor"' not in html  # env=remove suppresses anchors
