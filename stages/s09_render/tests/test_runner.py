from pathlib import Path

import yaml
from docx import Document
from PIL import Image

from stages.s09_render.runner import run


def test_run_produces_bundle_and_preview(tmp_path: Path):
    compose_dir = tmp_path / "compose"; compose_dir.mkdir()
    (compose_dir / "chapters").mkdir()
    (compose_dir / "chapters" / "01-intro.md").write_text(
        "# 1 引言\n\n这是引言。\n", encoding="utf-8"
    )
    (compose_dir / "chapters" / "02-results.md").write_text(
        "# 2 结果\n\n这是结果。\n", encoding="utf-8"
    )

    fig_dir = tmp_path / "fig"; fig_dir.mkdir()
    img_path = tmp_path / "imgs" / "a.jpg"; img_path.parent.mkdir()
    Image.new("RGB", (200, 100), "blue").save(img_path)
    (fig_dir / "fig_notes.yaml").write_text(yaml.safe_dump([
        {"fig_id": "Fig. 1", "image_abs_path": str(img_path),
         "caption": "图1: 测试",
         "deep_observation": "观察",
         "text_claim_check": [{"claim": "x", "verdict": "supported", "note": "ok"}]},
    ], allow_unicode=True), encoding="utf-8")

    out_dir = tmp_path / "out"
    run(compose_dir=compose_dir, fig_notes_dir=fig_dir, out_dir=out_dir,
        paper_title="测试论文")

    bundle = out_dir / "mypaper_bundle"
    assert (bundle / "chapters" / "01-intro.md").exists()
    assert (bundle / "chapters" / "02-results.md").exists()
    assert (bundle / "figures" / "a.jpg").exists()
    assert (bundle / "README.md").exists()

    preview = out_dir / "preview.docx"
    assert preview.exists() and preview.stat().st_size > 4000
    d = Document(preview)
    text = "\n".join(p.text for p in d.paragraphs)
    assert "引言" in text and "结果" in text


def test_run_clears_stale_bundle_files(tmp_path: Path):
    """Bundle from a previous run shouldn't bleed into a new run."""
    import yaml as _yaml
    compose_dir = tmp_path / "compose"; (compose_dir / "chapters").mkdir(parents=True)
    (compose_dir / "chapters" / "01-only.md").write_text("# 1 Only\n\nnew body\n", encoding="utf-8")

    fig_dir = tmp_path / "fig"; fig_dir.mkdir()
    (fig_dir / "fig_notes.yaml").write_text("[]", encoding="utf-8")

    out_dir = tmp_path / "out"
    bundle_chapters = out_dir / "mypaper_bundle" / "chapters"
    bundle_chapters.mkdir(parents=True)
    # Plant a stale file from a previous run
    (bundle_chapters / "99-stale.md").write_text("stale", encoding="utf-8")

    run(compose_dir=compose_dir, fig_notes_dir=fig_dir, out_dir=out_dir, paper_title="t")

    remaining = sorted(p.name for p in bundle_chapters.glob("*.md"))
    assert remaining == ["01-only.md"], remaining


def test_dedup_figure_embeds(tmp_path: Path):
    """Same Fig. N referenced in two chapters should be embedded only once."""
    import yaml as _yaml
    compose = tmp_path / "compose"; (compose / "chapters").mkdir(parents=True)
    (compose / "chapters" / "01-a.md").write_text("# 1 First\n\nbody mentions Fig. 1.\n", encoding="utf-8")
    (compose / "chapters" / "02-b.md").write_text("# 2 Second\n\nbody also mentions Fig. 1 here.\n", encoding="utf-8")

    fig_dir = tmp_path / "fig"; fig_dir.mkdir()
    img = tmp_path / "imgs" / "a.jpg"; img.parent.mkdir()
    Image.new("RGB", (300, 200), "red").save(img)
    (fig_dir / "fig_notes.yaml").write_text(
        _yaml.safe_dump([{"fig_id": "Fig. 1", "image_abs_path": str(img),
                          "caption": "test", "deep_observation": "obs"}],
                        allow_unicode=True),
        encoding="utf-8",
    )
    out_dir = tmp_path / "out"
    run(compose_dir=compose, fig_notes_dir=fig_dir, out_dir=out_dir, paper_title="t")
    d = Document(out_dir / "preview.docx")
    assert len(d.inline_shapes) == 1, f"expected 1 unique embed, got {len(d.inline_shapes)}"


def test_multi_panel_embed(tmp_path: Path):
    """A fig_note with image_paths=[a, b] embeds both images under one caption."""
    import yaml as _yaml
    from PIL import Image
    compose = tmp_path / "compose"; (compose / "chapters").mkdir(parents=True)
    (compose / "chapters" / "01-a.md").write_text("# 1 First\n\nmentions Fig. 1.\n", encoding="utf-8")

    img_a = tmp_path / "imgs" / "a.jpg"; img_a.parent.mkdir()
    img_b = tmp_path / "imgs" / "b.jpg"
    Image.new("RGB", (200, 100), "red").save(img_a)
    Image.new("RGB", (200, 100), "blue").save(img_b)
    fig_dir = tmp_path / "fig"; fig_dir.mkdir()
    (fig_dir / "fig_notes.yaml").write_text(
        _yaml.safe_dump([{"fig_id": "Fig. 1",
                          "image_paths": [str(img_a), str(img_b)],
                          "image_abs_path": str(img_a),
                          "caption": "two-panel",
                          "deep_observation": "obs"}],
                        allow_unicode=True),
        encoding="utf-8",
    )
    out_dir = tmp_path / "out"
    run(compose_dir=compose, fig_notes_dir=fig_dir, out_dir=out_dir, paper_title="t")
    from docx import Document
    d = Document(out_dir / "preview.docx")
    # Both panels embedded
    assert len(d.inline_shapes) == 2, f"expected 2 embeds, got {len(d.inline_shapes)}"


def test_run_uses_title_from_context_yaml(tmp_path: Path):
    """When context_dir contains a context.yaml with `title:`, use it."""
    import yaml as _y
    compose = tmp_path / "compose"; (compose / "chapters").mkdir(parents=True)
    (compose / "chapters" / "01.md").write_text("# C\n\nbody\n", encoding="utf-8")

    fig_dir = tmp_path / "fig"; fig_dir.mkdir()
    (fig_dir / "fig_notes.yaml").write_text("[]", encoding="utf-8")

    context_dir = tmp_path / "ctx"; context_dir.mkdir()
    (context_dir / "context.yaml").write_text(
        _y.safe_dump({"title": "Real Paper Title From OCR"}, allow_unicode=True),
        encoding="utf-8",
    )

    out_dir = tmp_path / "out"
    from stages.s09_render.runner import run
    run(compose_dir=compose, fig_notes_dir=fig_dir, context_dir=context_dir,
        out_dir=out_dir, paper_title="fallback-id", lang="en", formats=["docx"])

    d = Document(out_dir / "preview.docx")
    text = "\n".join(p.text for p in d.paragraphs)
    assert "Real Paper Title From OCR" in text
    assert "fallback-id" not in text


def test_run_falls_back_when_context_yaml_missing(tmp_path: Path):
    """If context_dir doesn't exist or context.yaml is absent, use paper_title."""
    compose = tmp_path / "compose"; (compose / "chapters").mkdir(parents=True)
    (compose / "chapters" / "01.md").write_text("# C\n\nbody\n", encoding="utf-8")

    fig_dir = tmp_path / "fig"; fig_dir.mkdir()
    (fig_dir / "fig_notes.yaml").write_text("[]", encoding="utf-8")

    out_dir = tmp_path / "out"
    from stages.s09_render.runner import run
    run(compose_dir=compose, fig_notes_dir=fig_dir, context_dir=tmp_path / "missing",
        out_dir=out_dir, paper_title="fallback-id", lang="en", formats=["docx"])

    d = Document(out_dir / "preview.docx")
    text = "\n".join(p.text for p in d.paragraphs)
    assert "fallback-id" in text
