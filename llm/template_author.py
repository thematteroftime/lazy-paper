"""Template author (v1.15): user idea -> question-driven outline docx.

Drafts an s05-compatible outline docx for ONE paper: cheap prescan (pdfplumber
or existing run artifacts) + one text-LLM call + deterministic docx writer.
The generated docx round-trips losslessly through s05's parser because only
manually numbered lines ("1 Title") become headings there — every question is
a plain paragraph and lands in `guidance` (s05 treats "?"-ending lines as
guidance by rule).
"""
from __future__ import annotations

import json
import re
from pathlib import Path

import pdfplumber
import yaml

from llm.client import LLM
from stages._common import load_yaml, safe_parse_yaml

PROMPT_PATH = Path(__file__).resolve().parent / "prompts" / "template_author.md"

# Mirror of s05's _NUMBERED_RE (stages/s05_template/runner.py): lines matching
# this would be promoted to headings, so questions must never match it.
_NUMBERED_RE = re.compile(r"^\s*(\d+(?:\.\d+){0,2})\s+(.+?)\s*$")
_LEADING_NUM_RE = re.compile(r"^\s*\d+(?:\.\d+){0,2}[.)]?\s+")
_FENCE_RE = re.compile(r"^\s*```[a-zA-Z]*\s*|\s*```\s*$")


def prescan_pdf(pdf: Path, *, max_pages: int = 4, max_chars: int = 6000) -> str:
    """Cheap text-layer extraction (no OCR API). Good enough for title/abstract."""
    parts: list[str] = []
    with pdfplumber.open(pdf) as doc:
        for page in doc.pages[:max_pages]:
            parts.append(page.extract_text() or "")
    text = "\n".join(p for p in parts if p).strip()
    if not text:
        raise SystemExit(
            f"prescan: no text layer in {pdf} — run s01 OCR first and use --run instead")
    return text[:max_chars]


def prescan_run(run_dir: Path, *, max_chars: int = 6000) -> str:
    """Digest from existing run artifacts: context > chapters > captions > OCR head."""
    parts: list[str] = []
    ctx_path = run_dir / "s06_context" / "context.yaml"
    if ctx_path.exists():
        ctx = load_yaml(ctx_path) or {}
        parts.append(yaml.safe_dump(ctx, allow_unicode=True, sort_keys=False))
    idx_path = run_dir / "s03_chapter" / "chapter_index.yaml"
    if idx_path.exists():
        # s03 writes a top-level LIST of {chapter_no, title, file, ...};
        # tolerate a {chapters: [...]} wrapper for forward compatibility.
        idx = load_yaml(idx_path) or []
        chapters = (idx.get("chapters") or []) if isinstance(idx, dict) else idx
        titles = [str(c.get("title", "")) for c in chapters
                  if isinstance(c, dict) and c.get("title")]
        if titles:
            parts.append("Chapter titles: " + " | ".join(titles))
    figs_path = run_dir / "s04_figures" / "figures.yaml"
    if figs_path.exists():
        figs = load_yaml(figs_path) or []
        captions = [f"{f.get('fig_id', '')}: {f.get('caption', '')}"
                    for f in figs if f.get("caption")][:15]
        if captions:
            parts.append("Figure captions:\n" + "\n".join(captions))
    docs = sorted((run_dir / "s02_clean").glob("doc_*.md"))
    if docs:
        parts.append(docs[0].read_text(encoding="utf-8")[:2500])
    digest = "\n\n".join(p for p in parts if p.strip()).strip()
    if not digest:
        raise SystemExit(
            f"prescan: no usable artifacts under {run_dir} "
            f"(need any of s06 context.yaml / s03 chapter_index.yaml / "
            f"s04 figures.yaml / s02 doc_*.md)")
    return digest[:max_chars]


_LANG_INSTRUCTIONS = {
    "zh": "All titles and questions in Chinese (keep established English technical terms as-is).",
    "en": "All titles and questions in English.",
}


def _split_prompt(template_text: str) -> tuple[str, str]:
    system_marker, user_marker = "SYSTEM:", "USER:"
    sys_start = template_text.index(system_marker) + len(system_marker)
    user_start = template_text.index(user_marker)
    return (template_text[sys_start:user_start].strip(),
            template_text[user_start + len(user_marker):].strip())


def _clean_title(t: str) -> str:
    t = _LEADING_NUM_RE.sub("", str(t).strip())
    t = t.rstrip("?？ 。.").strip()
    t = re.sub(r"^[^0-9A-Za-z一-鿿]+", "", t).strip()  # no '(' '-' '·' starts
    if t and t[0].islower():
        t = t[0].upper() + t[1:]                               # no lowercase starts
    return t[:80] or "Untitled"


def _clean_question(q: str) -> str:
    q = str(q).strip()
    # A question starting like "3 个指标…" would match s05's numbered-heading
    # regex; the "- " guidance prefix makes it un-promotable by construction.
    if _NUMBERED_RE.match(q):
        q = "- " + q
    return q


def _validate(data) -> list[dict]:
    secs = data.get("sections") if isinstance(data, dict) else None
    out: list[dict] = []
    for s in secs or []:
        if not isinstance(s, dict) or not s.get("title"):
            continue
        qs = [_clean_question(q) for q in (s.get("questions") or []) if str(q).strip()]
        if qs:
            out.append({"title": _clean_title(s["title"]), "questions": qs})
    if not out:
        raise SystemExit(
            "template draft: LLM returned no usable sections "
            "(see the saved .response.json next to the output path)")
    return out


def draft(*, idea: str, paper_digest: str, library_context: str = "",
          lang: str = "zh", n_sections: int = 6,
          audit_base: Path | None = None):
    """Text-LLM call (1 retry) -> validated [{title, questions}] outline.

    Returns (sections, LLMResponse). When `audit_base` is given, the prompt
    and raw response are persisted BEFORE validation, so a rejected response
    is always inspectable at <audit_base>.response.json.
    """
    system_tpl, user_tpl = _split_prompt(PROMPT_PATH.read_text(encoding="utf-8"))
    system = system_tpl.format(
        n_sections=n_sections,
        lang_instruction=_LANG_INSTRUCTIONS.get(lang, _LANG_INSTRUCTIONS["en"]),
    )
    user = user_tpl.format(idea=idea, paper_digest=paper_digest,
                           library_context=library_context or "(empty)")
    llm = LLM(role="text")
    last_err: SystemExit | None = None
    for attempt in range(2):
        attempt_user = user if attempt == 0 else (
            user + "\n\nYour previous output was not valid YAML for the "
                   "required schema. Output ONLY the YAML — no fence, "
                   "no preamble, no commentary.")
        resp = llm.chat(system=system, user=attempt_user,
                        temperature=0.4, max_tokens=4096)
        if audit_base is not None:
            base = Path(audit_base)
            base.parent.mkdir(parents=True, exist_ok=True)
            Path(str(base) + ".prompt.md").write_text(
                f"SYSTEM:\n{system}\n\nUSER:\n{attempt_user}", encoding="utf-8")
            Path(str(base) + ".response.json").write_text(
                json.dumps({"model": resp.model, "usage": resp.usage,
                            "content": resp.content},
                           ensure_ascii=False, indent=2),
                encoding="utf-8")
        cleaned = _FENCE_RE.sub("", resp.content.strip())
        try:
            return _validate(safe_parse_yaml(cleaned)), resp
        except SystemExit as e:
            last_err = e
    raise last_err


def write_docx(sections: list[dict], out_path: Path, *, idea: str) -> None:
    """Plain-paragraph docx that round-trips deterministically through s05.

    Heading lines are manually numbered ("1 Title") — the ONLY pattern s05
    promotes. Questions are plain paragraphs -> guidance. The preamble line
    starts with "(" so s05 ignores it (guidance prefix, no open heading).
    """
    from docx import Document  # heavy import; keep local like s09 does

    doc = Document()
    doc.add_paragraph(f"(auto-generated by lazy-paper template — idea: {idea[:120]})")
    for i, sec in enumerate(sections, 1):
        doc.add_paragraph(f"{i} {sec['title']}")
        for q in sec["questions"]:
            doc.add_paragraph(q)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def roundtrip_check(out_path: Path, sections: list[dict]) -> list[dict]:
    """Parse the written docx with the REAL s05 parser; hard-fail on drift."""
    from stages.s05_template.runner import parse_template

    nodes = parse_template(out_path)
    if len(nodes) != len(sections):
        raise SystemExit(
            f"template self-check failed: wrote {len(sections)} sections but "
            f"s05 parsed {len(nodes)} — please report this as a bug")
    return nodes


def library_context(lib, idea: str, *, top_k: int = 5) -> str:
    """Manifest summary + idea-relevant excerpts for the drafting prompt."""
    manifest = lib.papers()
    if not manifest:
        return ""
    lines = ["Library papers:"]
    for pid, e in manifest.items():
        kw = ", ".join((e.get("keywords") or [])[:5])
        lines.append(f"- {pid}: {e.get('title', '')} (keywords: {kw})")
    hits = lib.query(idea, top_k=top_k)
    if hits:
        lines.append("")
        lines.append("Idea-relevant excerpts:")
        for h in hits:
            snippet = " ".join(h["text"].split())[:300]
            lines.append(f"- [{h['paper_id']}] {snippet}")
    return "\n".join(lines)
